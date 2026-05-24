from flask import Blueprint, request, Response, send_file
from openai import OpenAI
from dotenv import load_dotenv
import os, json, io, pdfplumber
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
press_bp = Blueprint("press", __name__)

def load_prompt():
    path = os.path.join(os.path.dirname(__file__), "../prompts/press_prompt.txt")
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def respond(data, code=200):
    return Response(
        json.dumps(data, ensure_ascii=False),
        status=code,
        mimetype="application/json; charset=utf-8"
    )

def extract_pdf_text(file_bytes):
    try:
        text = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except:
        return ""

def make_docx(text):
    """보도자료 텍스트를 Word 문서로 변환"""
    doc = Document()

    # 페이지 여백
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

    # 기본 폰트
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)

    # 상단 기관명
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run("한성대학교 보도자료")
    header_run.bold = True
    header_run.font.size = Pt(11)
    header_run.font.color.rgb = RGBColor(0, 32, 96)

    # 구분선
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_p.add_run("─" * 44)
    line_run.font.color.rgb = RGBColor(0, 32, 96)
    line_run.font.size = Pt(9)

    doc.add_paragraph()

    lines = text.strip().split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            doc.add_paragraph()
            i += 1
            continue

        # [제목]
        if line.startswith("[제목]"):
            i += 1
            if i < len(lines) and lines[i].strip():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(lines[i].strip())
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0, 32, 96)
                i += 1
            continue

        # [소제목]
        if line.startswith("[소제목]"):
            i += 1
            while i < len(lines):
                sub = lines[i].strip()
                if not sub or (not sub.startswith("-") and sub.startswith("[")):
                    break
                if sub.startswith("-"):
                    sub = sub.lstrip("- ").strip()
                if sub:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(sub)
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 61, 165)
                i += 1
            doc.add_paragraph()
            continue

        # [본문] — 섹션 헤더만 스킵, 내용은 그대로
        if line == "[본문]":
            i += 1
            continue

        # [문의]
        if line.startswith("[문의]"):
            doc.add_paragraph()
            border_p = doc.add_paragraph()
            border_run = border_p.add_run("─" * 44)
            border_run.font.color.rgb = RGBColor(0, 32, 96)
            border_run.font.size = Pt(9)

            i += 1
            while i < len(lines):
                contact = lines[i].strip()
                if contact:
                    p = doc.add_paragraph()
                    run = p.add_run(contact)
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(100, 100, 100)
                i += 1
            break

        # 일반 본문
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Pt(20)
        run = p.add_run(line)
        run.font.size = Pt(10.5)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

@press_bp.route("/press", methods=["POST"])
def generate_press():
    # PDF 업로드 방식
    if "pdf" in request.files:
        pdf_file   = request.files["pdf"]
        dept_name  = request.form.get("dept_name", "")
        press_type = request.form.get("press_type", "")
        summary    = request.form.get("summary", "")
        pdf_text   = extract_pdf_text(pdf_file.read())

        if pdf_text:
            user_content = f"""
아래 PDF 문서 내용을 바탕으로 한성대학교 보도자료를 작성해주세요.
부서명: {dept_name}
행사·사업 유형: {press_type}
추가 요약: {summary}
PDF 내용:
{pdf_text[:3000]}
"""
        else:
            return respond({"error": "PDF에서 텍스트를 추출할 수 없습니다."}, 400)
    else:
        data       = request.json or {}
        dept_name  = data.get("dept_name", "")
        press_type = data.get("press_type", "")
        summary    = data.get("summary", "")
        if not summary:
            return respond({"error": "내용을 입력해주세요."}, 400)
        user_content = f"""
부서명: {dept_name}
행사·사업 유형: {press_type}
핵심 내용: {summary}
"""

    system_prompt = load_prompt()
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_content},
        ],
        max_tokens=2000,
    )

    result = response.choices[0].message.content
    return respond({"result": result})


@press_bp.route("/press/download", methods=["POST"])
def download_press():
    """생성된 보도자료를 Word 파일로 다운로드"""
    data = request.json or {}
    text = data.get("text", "")
    if not text:
        return respond({"error": "내용이 없습니다."}, 400)

    buf = make_docx(text)
    return send_file(
        buf,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name="한성대학교_보도자료.docx"
    )
